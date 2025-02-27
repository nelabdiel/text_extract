import fitz  # PyMuPDF
import pdfplumber
import pytesseract
import io
import pandas as pd
import docx
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PIL import Image

app = Flask(__name__)
CORS(app)

def extract_text_from_pdf(pdf_bytes):
    """Extracts selectable text from a PDF loaded in memory."""
    # Open PDF from memory
    doc = fitz.open("pdf", pdf_bytes.getvalue())
    return "\n".join([page.get_text("text") for page in doc])

def extract_text_from_docx(docx_bytes):
    """Extracts text from a DOCX file loaded in memory."""
    # Load as in-memory stream
    doc_stream = io.BytesIO(docx_bytes.getvalue())
    doc = docx.Document(doc_stream)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_images_from_pdf(pdf_bytes):
    """Extracts images from a PDF and processes them in-memory without saving."""
    # Open from memory
    doc = fitz.open("pdf", pdf_bytes.getvalue())
    extracted_images = []

    for page_number, page in enumerate(doc, start=1):
        for img_index, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]  
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            # Convert to OCR-readable format
            img_stream = io.BytesIO(image_bytes)
            img = Image.open(img_stream)

            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format="PNG")
            # Store bytes in memory
            extracted_images.append(img_byte_arr.getvalue())

    return extracted_images

def extract_images_from_docx(docx_bytes):
    """Extracts images from a DOCX file and processes them in-memory."""
    doc_stream = io.BytesIO(docx_bytes.getvalue())  
    doc = docx.Document(doc_stream)
    extracted_images = []

    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            image = doc.part.rels[rel].target_part.blob
            img = Image.open(io.BytesIO(image))

            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format="PNG")
            # Store bytes in memory
            extracted_images.append(img_byte_arr.getvalue())

    return extracted_images

def ocr_images(image_data_list):
    """Runs OCR on in-memory images and returns extracted text."""
    ocr_text = ""
    for image_data in image_data_list:
        img = Image.open(io.BytesIO(image_data))
        ocr_text += pytesseract.image_to_string(img) + "\n"
    return ocr_text

def extract_tables_from_pdf(pdf_bytes):
    """Extracts tables from a PDF and returns them as JSON."""
    extracted_tables = []
    
    with pdfplumber.open(io.BytesIO(pdf_bytes.getvalue())) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()
            for table_index, table in enumerate(tables, start=1):
                df = pd.DataFrame(table)  
                df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')
                
                extracted_tables.append({
                    "page": page_number,
                    "table_index": table_index,
                    # Convert to JSON
                    "data": df.to_dict(orient="records")
                })

    return extracted_tables

def extract_tables_from_docx(docx_bytes):
    """Extracts tables from a DOCX file and returns them as JSON."""
    doc_stream = io.BytesIO(docx_bytes.getvalue())  
    doc = docx.Document(doc_stream)
    extracted_tables = []

    for table_index, table in enumerate(doc.tables, start=1):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        extracted_tables.append({
            "table_index": table_index,
            "data": table_data
        })

    return extracted_tables

@app.route("/extract", methods=["POST"])
def extract():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    # Get file extension
    file_ext = filename.split(".")[-1].lower()

    # Load file into memory
    file_bytes = io.BytesIO(file.read())

    if file_ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
        image_data_list = extract_images_from_pdf(file_bytes)
        ocr_text = ocr_images(image_data_list) if image_data_list else ""
        tables = extract_tables_from_pdf(file_bytes)
    
    elif file_ext == "docx":
        text = extract_text_from_docx(file_bytes)
        image_data_list = extract_images_from_docx(file_bytes)
        ocr_text = ocr_images(image_data_list) if image_data_list else ""
        tables = extract_tables_from_docx(file_bytes)

    else:
        return jsonify({"error": "Unsupported file type. Please upload a PDF or DOCX."}), 400

    return jsonify({
        "message": "Extraction complete",
        "file_type": file_ext.upper(),
        "extracted_text": text,
        "ocr_text": ocr_text,
        "extracted_images": len(image_data_list),  
        "extracted_tables": tables
    })

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5005, debug=True)
